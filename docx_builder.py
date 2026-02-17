"""DocxBuilder — построитель DOCX из HTML с сохранением пустых строк"""

import re
import logging
from dataclasses import dataclass
from typing import Callable

from bs4 import NavigableString, Tag
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

from config import HAS_PILLOW, AppConfig
from utils import add_horizontal_line, add_hyperlink
from network import download_image, ImageCache

if HAS_PILLOW:
    from PIL import Image

logger = logging.getLogger(__name__)


@dataclass
class StyleContext:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    spoiler: bool = False
    code: bool = False

    def copy(self) -> 'StyleContext':
        return StyleContext(
            bold=self.bold, italic=self.italic, underline=self.underline,
            strike=self.strike, spoiler=self.spoiler, code=self.code,
        )


class DocxBuilder:
    BLOCK_TAGS = frozenset([
        'div', 'p', 'blockquote', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'table', 'ul', 'ol', 'hr', 'pre',
    ])
    BOLD_TAGS = frozenset(['b', 'strong'])
    ITALIC_TAGS = frozenset(['i', 'em'])
    UNDERLINE_TAGS = frozenset(['u', 'ins'])
    STRIKE_TAGS = frozenset(['s', 'strike', 'del'])
    CODE_TAGS = frozenset(['code', 'pre'])
    HEADING_CLASSES = {'bb_h1': 1, 'bb_h2': 2, 'bb_h3': 3}

    MAX_RECURSION_DEPTH = 50
    MAX_LIST_DEPTH = 10

    def __init__(self, doc_context, config=None, session=None,
                 image_cache=None, log_func=None):
        self.doc = doc_context
        self.config = config or AppConfig()
        self.session = session
        self.image_cache = image_cache
        self.log_func = log_func or (lambda msg: None)
        self.current_paragraph = None
        self.is_cell = not hasattr(self.doc, 'add_heading')
        self._depth = 0
        self._list_depth = 0

        # === Трекер пустых строк ===
        # Считает последовательные <br> для создания пустых абзацев
        self._consecutive_br = 0
        # Флаг: был ли уже текст/контент (чтобы не ставить
        # пустые строки в самом начале)
        self._has_content = False
        # Флаг: текущий параграф пуст (только создан, без текста)
        self._paragraph_is_empty = True

    def get_paragraph(self, style=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph(style=style)
            self.current_paragraph.alignment = alignment
            pf = self.current_paragraph.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            self._paragraph_is_empty = True
        return self.current_paragraph

    def close_paragraph(self):
        """Завершить текущий параграф"""
        if self.current_paragraph is not None:
            self._has_content = True
        self.current_paragraph = None
        self._paragraph_is_empty = True

    def _add_empty_paragraph(self):
        """Добавить пустой параграф (визуальная пустая строка)"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # Добавляем пустой run чтобы параграф не схлопнулся
        run = p.add_run("")
        run.font.size = Pt(11)

    def _flush_pending_breaks(self):
        """
        Обработать накопленные <br>.
        
        Логика:
        - 1 <br> = просто новый параграф (перенос строки)
        - 2+ <br> = пустые строки между абзацами
        
        Вызывается ПЕРЕД добавлением любого контента.
        """
        if self._consecutive_br <= 0:
            return

        if not self._has_content:
            # В самом начале документа не ставим пустые строки
            self._consecutive_br = 0
            return

        # Первый <br> — это просто перенос (новый параграф)
        # Каждый следующий — пустая строка
        empty_lines = self._consecutive_br - 1

        for _ in range(empty_lines):
            self._add_empty_paragraph()

        self._consecutive_br = 0

    def _apply_style(self, run, ctx):
        if ctx.bold: run.bold = True
        if ctx.italic: run.italic = True
        if ctx.underline: run.underline = True
        if ctx.strike: run.font.strike = True
        if ctx.spoiler:
            run.font.highlight_color = WD_COLOR_INDEX.BLACK
            run.font.color.rgb = RGBColor(255, 255, 255)
        if ctx.code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    def _update_context(self, node, ctx):
        new_ctx = ctx.copy()
        classes = set(node.get('class', []))
        if node.name in self.BOLD_TAGS: new_ctx.bold = True
        if node.name in self.ITALIC_TAGS: new_ctx.italic = True
        if node.name in self.UNDERLINE_TAGS: new_ctx.underline = True
        if node.name in self.STRIKE_TAGS or 'bb_strike' in classes:
            new_ctx.strike = True
        if 'bb_spoiler' in classes: new_ctx.spoiler = True
        if node.name in self.CODE_TAGS or 'bb_code' in classes:
            new_ctx.code = True
        for cls_name in self.HEADING_CLASSES:
            if cls_name in classes:
                new_ctx.bold = True
                break
        return new_ctx

    def _add_image(self, src):
        self._flush_pending_breaks()
        self.close_paragraph()
        img_data = download_image(
            src, session=self.session,
            config=self.config, cache=self.image_cache
        )
        if not img_data:
            return
        try:
            max_w = (self.config.cell_image_width_inches
                     if self.is_cell
                     else self.config.max_image_width_inches)
            final_width = Inches(max_w)
            if HAS_PILLOW:
                try:
                    img = Image.open(img_data)
                    width_px, _ = img.size
                    img_data.seek(0)
                    if width_px < 400:
                        final_width = Inches(width_px / 96.0)
                except Exception:
                    img_data.seek(0)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(img_data, width=final_width)
            self._has_content = True
        except Exception as e:
            logger.warning(f"Ошибка вставки изображения: {e}")
        self.close_paragraph()

    # ==========================================
    # ГЛАВНЫЙ ОБРАБОТЧИК
    # ==========================================

    def process_node(self, node, style_ctx=None):
        self._depth += 1
        if self._depth > self.MAX_RECURSION_DEPTH:
            self._depth -= 1
            return
        try:
            if style_ctx is None:
                style_ctx = StyleContext()
            if isinstance(node, NavigableString):
                self._process_text(node, style_ctx)
            elif isinstance(node, Tag):
                self._process_tag(node, style_ctx)
        finally:
            self._depth -= 1

    def _process_text(self, node, ctx):
        text = str(node)

        if not ctx.code:
            text = re.sub(r'\s+', ' ', text)

        if not text or (text.isspace() and not ctx.code):
            return

        # Есть реальный текст — сбрасываем накопленные переносы
        self._flush_pending_breaks()

        p = self.get_paragraph()

        # Убираем ведущие пробелы в начале параграфа
        if self._paragraph_is_empty:
            text = text.lstrip()
            if not text:
                return

        run = p.add_run(text)
        self._apply_style(run, ctx)
        self._paragraph_is_empty = False
        self._has_content = True

    def _process_tag(self, node, ctx):
        tag = node.name
        classes = set(node.get('class', []))

        # === <br> — специальная обработка ===
        if tag == 'br':
            self._consecutive_br += 1
            self.close_paragraph()
            return

        # === Блочные теги ===
        if tag in self.BLOCK_TAGS:
            # Проверяем: пустой ли блок (содержит только пробелы/br)
            if self._is_empty_block(node):
                # Пустой <div> или <p> = пустая строка
                if self._has_content:
                    self._consecutive_br += 1
                    self.close_paragraph()
                return

            # Непустой блочный тег — сбрасываем переносы
            self._flush_pending_breaks()
            self.close_paragraph()

        new_ctx = self._update_context(node, ctx)

        # --- Специальные обработчики ---

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self._flush_pending_breaks()
            self._handle_heading(node, tag)
            return

        for cls_name, level in self.HEADING_CLASSES.items():
            if cls_name in classes:
                self._flush_pending_breaks()
                self._handle_steam_heading(node, level)
                return

        if tag == 'hr':
            self._flush_pending_breaks()
            p = self.doc.add_paragraph()
            add_horizontal_line(p)
            self.close_paragraph()
            self._has_content = True
            return

        if tag == 'img':
            src = node.get('src')
            if src:
                self._add_image(src)
            return

        if tag == 'a':
            self._flush_pending_breaks()
            self._handle_link(node, new_ctx)
            return

        if tag in ('ul', 'ol'):
            self._flush_pending_breaks()
            self._handle_list(node, tag, new_ctx)
            return

        if tag == 'blockquote':
            self._flush_pending_breaks()
            self._handle_blockquote(node, new_ctx)
            return

        if 'bb_table' in classes:
            self._flush_pending_breaks()
            self._handle_table(node)
            return

        # --- Рекурсия ---
        for child in node.children:
            self.process_node(child, new_ctx)

        if new_ctx.code and tag in self.CODE_TAGS:
            self.close_paragraph()

    # ==========================================
    # ПРОВЕРКА ПУСТОГО БЛОКА
    # ==========================================

    def _is_empty_block(self, node):
        """
        Проверяет, содержит ли блочный элемент только
        пробелы, <br> или ничего.
        
        Примеры пустых блоков:
          <div></div>
          <p></p>
          <div><br></div>
          <div>   </div>
          <p>&nbsp;</p>
        """
        for child in node.children:
            if isinstance(child, Tag):
                if child.name == 'br':
                    continue
                # Есть непустой дочерний тег — блок не пустой
                return False
            elif isinstance(child, NavigableString):
                text = str(child).strip()
                # &nbsp; тоже считаем пустым
                text = text.replace('\xa0', '').replace('&nbsp;', '')
                if text:
                    return False
        return True

    # ==========================================
    # ОБРАБОТЧИКИ ТЕГОВ
    # ==========================================

    def _handle_heading(self, node, tag_name):
        text = node.get_text(strip=True)
        if not text:
            self.close_paragraph()
            return
        try:
            level = int(tag_name[1])
        except (ValueError, IndexError):
            level = 1
        if not self.is_cell:
            self.doc.add_heading(text, level=min(level, 9))
        else:
            p = self.get_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(max(14 - level, 9))
        self.close_paragraph()
        self._has_content = True

    def _handle_steam_heading(self, node, level):
        text = node.get_text(strip=True)
        if not text:
            self.close_paragraph()
            return
        if not self.is_cell:
            self.doc.add_heading(text, level=level + 1)
        else:
            p = self.get_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
        self.close_paragraph()
        self._has_content = True

    def _handle_link(self, node, ctx):
        href = node.get('href')
        img_child = node.find('img')
        if img_child:
            src = img_child.get('src')
            if src:
                self._add_image(src)
            return

        p = self.get_paragraph()
        link_text = node.get_text().strip()

        if not link_text:
            for child in node.children:
                self.process_node(child, ctx)
            return

        if href:
            add_hyperlink(p, href, link_text)
        else:
            run = p.add_run(link_text)
            self._apply_style(run, ctx)

        self._paragraph_is_empty = False
        self._has_content = True

    def _handle_list(self, node, list_type, ctx):
        self._list_depth += 1
        if self._list_depth > self.MAX_LIST_DEPTH:
            self._list_depth -= 1
            return
        try:
            style = 'List Number' if list_type == 'ol' else 'List Bullet'
            for li in node.find_all('li', recursive=False):
                self.current_paragraph = self.doc.add_paragraph(
                    style=style
                )
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = self.current_paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                if self._list_depth > 1:
                    pf.left_indent = Inches(0.25 * self._list_depth)
                self._paragraph_is_empty = True
                for child in li.children:
                    self.process_node(child, ctx)
                self.close_paragraph()
            self._has_content = True
        finally:
            self._list_depth -= 1

    def _handle_blockquote(self, node, ctx):
        self.close_paragraph()
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        quote_ctx = ctx.copy()
        quote_ctx.italic = True
        self.current_paragraph = p
        self._paragraph_is_empty = True
        for child in node.children:
            self.process_node(child, quote_ctx)
        self.close_paragraph()
        self._has_content = True

    def _handle_table(self, table_node):
        self.close_paragraph()
        if self.is_cell:
            p = self.get_paragraph()
            p.add_run("[Table]").italic = True
            self.close_paragraph()
            return
        rows = table_node.find_all('div', class_='bb_table_tr')
        if not rows:
            return
        first_cells = rows[0].find_all(
            'div', class_=['bb_table_td', 'bb_table_th']
        )
        cols = len(first_cells)
        if cols == 0:
            return
        try:
            table = self.doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(
                    'div', class_=['bb_table_td', 'bb_table_th']
                )
                for j, cell_html in enumerate(cells):
                    if j >= cols:
                        break
                    cell_docx = table.rows[i].cells[j]
                    cell_docx._element.clear_content()
                    cb = DocxBuilder(
                        cell_docx, config=self.config,
                        session=self.session,
                        image_cache=self.image_cache,
                        log_func=self.log_func
                    )
                    for child in cell_html.children:
                        cb.process_node(child)
                    if len(cell_docx.paragraphs) == 0:
                        p = cell_docx.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                    else:
                        for p in cell_docx.paragraphs:
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
            self._has_content = True
        except Exception as e:
            logger.error(f"Ошибка таблицы: {e}")
        self.close_paragraph()