"""Парсер руководств Steam — обновлённая секция PDF"""

import os
import re
import logging
import threading
from typing import Callable

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

from config import AppConfig
from translations import get_text
from utils import clean_filename
from network import create_session, URLValidator, ImageCache
from docx_builder import DocxBuilder
from pdf_converter import convert_docx_to_pdf, check_available_converters

logger = logging.getLogger(__name__)


class GuideDownloader:
    def __init__(self, config: AppConfig):
        self.config = config
        self.session = create_session(config)
        self.image_cache = ImageCache(max_size=100)
        self._cancelled = threading.Event()

    def cancel(self):
        self._cancelled.set()

    @property
    def is_cancelled(self):
        return self._cancelled.is_set()

    def download(self, url, save_dir, lang_code, log_func,
                 finish_func, convert_pdf=False):
        self._cancelled.clear()
        self.image_cache.clear()
        try:
            self._do_download(url, save_dir, lang_code,
                              log_func, convert_pdf)
        except Exception as e:
            logger.error(f"Ошибка загрузки: {e}", exc_info=True)
            log_func(f"Error: {e}")
        finally:
            logger.debug(self.image_cache.stats)
            finish_func()

    def _do_download(self, url, save_dir, lang_code,
                     log_func, convert_pdf):
        T = lambda key, *a: get_text(lang_code, key, *a)
        log_func(T("log_start", url))

        # Предварительная проверка PDF-конвертера
        if convert_pdf:
            converters = check_available_converters()
            has_any = any(converters.values())
            if not has_any:
                log_func(f"⚠ {T('err_pdf_no_support')}")
                log_func("  Continuing with DOCX only...")
                convert_pdf = False

        if not os.path.exists(save_dir):
            try:
                os.makedirs(save_dir, exist_ok=True)
            except OSError as e:
                log_func(f"{T('err_creating_dir')} {e}")
                return

        if self.is_cancelled:
            log_func(T("log_cancelled"))
            return

        try:
            response = self.session.get(url, timeout=self.config.timeout)
            response.raise_for_status()
            response.encoding = 'utf-8'
        except requests.ConnectionError:
            log_func(T("err_net_connection"))
            return
        except requests.Timeout:
            log_func(T("err_net_timeout"))
            return
        except requests.HTTPError as e:
            log_func(T("err_access", e.response.status_code))
            return
        except requests.RequestException as e:
            log_func(f"{T('err_net')} {e}")
            return

        if self.is_cancelled:
            log_func(T("log_cancelled"))
            return

        soup = BeautifulSoup(response.text, 'html.parser')
        doc = Document()
        self._setup_styles(doc)

        guide_title = self._extract_title(soup)
        doc.add_heading(guide_title, 0)

        safe_title = clean_filename(guide_title)
        if not safe_title or len(safe_title) < 2:
            gid = URLValidator.extract_guide_id(url)
            safe_title = f"manual_{gid}" if gid else "manual_unknown"

        full_path = os.path.join(save_dir, f"{safe_title}.docx")
        log_func(T("log_file_target", full_path))

        if self.is_cancelled:
            log_func(T("log_cancelled"))
            return

        builder = DocxBuilder(
            doc, config=self.config, session=self.session,
            image_cache=self.image_cache, log_func=log_func
        )

        if not self._process_content(soup, doc, builder,
                                     lang_code, log_func):
            log_func(T("err_content"))
            return

        if self.is_cancelled:
            log_func(T("log_cancelled"))
            return

        # Сохранение DOCX
        try:
            doc.save(full_path)
            log_func(T("log_success", full_path))
        except PermissionError:
            log_func(T("err_permission"))
            return
        except OSError as e:
            log_func(f"Error: {e}")
            return

        # Конвертация в PDF
        if convert_pdf and not self.is_cancelled:
            log_func(T("log_pdf_converting"))
            success, result = convert_docx_to_pdf(full_path, log_func)
            if success:
                log_func(T("log_pdf_success", result))
            else:
                log_func(f"⚠ {T('err_pdf_failed')}")
                log_func(result)

    def _setup_styles(self, doc):
        try:
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
        except Exception as e:
            logger.warning(f"Ошибка стилей: {e}")

    def _extract_title(self, soup):
        title_node = soup.find('div', class_='workshopItemTitle')
        if title_node:
            t = title_node.get_text(strip=True)
            if t:
                return t
        title_tag = soup.find('title')
        if title_tag:
            t = title_tag.get_text(strip=True)
            t = re.sub(r'\s*::\s*Steam Community.*$', '', t)
            if t:
                return t
        return "Steam_Guide"

    def _process_content(self, soup, doc, builder,
                         lang_code, log_func):
        T = lambda key, *a: get_text(lang_code, key, *a)
        sections = soup.find_all('div', class_='subSection detailBox')

        if sections:
            log_func(T("log_sections_found", len(sections)))
            for section in sections:
                if self.is_cancelled:
                    return True
                title_div = section.find('div', class_='subSectionTitle')
                if title_div:
                    ch = title_div.get_text(" ", strip=True)
                    doc.add_heading(ch, 1)
                    short = ch[:40] + ("..." if len(ch) > 40 else "")
                    log_func(T("log_processing", short))
                desc_div = section.find('div', class_='subSectionDesc')
                if desc_div:
                    for child in desc_div.children:
                        if self.is_cancelled:
                            return True
                        builder.process_node(child)
                    builder.close_paragraph()
            return True

        content = (
            soup.find('div', id='guideContent')
            or soup.find('div', class_='guide subSections')
        )
        if content:
            log_func(T("log_processing", "main content"))
            for child in content.children:
                if self.is_cancelled:
                    return True
                builder.process_node(child)
            builder.close_paragraph()
            return True

        return False